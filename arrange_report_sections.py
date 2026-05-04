from pathlib import Path

from docx import Document
from docx.oxml.text.paragraph import CT_P


TEMP = Path(r"C:\Users\Nishant\AppData\Local\Temp\report_arrange_work.docx")
OUT = Path(r"C:\Users\Nishant\OneDrive\Desktop\report.docx")
COPY = Path(r"C:\Users\Nishant\OneDrive\Desktop\report_arranged_50_pages.docx")


RENAME = {
    "Extended Section 1: Detailed Problem Background": "Additional Project Documentation",
    "Extended Section 2: Dataset Description": "9.1 Dataset Description",
    "Extended Section 3: Data Cleaning Procedure": "9.2 Data Cleaning Procedure",
    "Extended Section 4: Feature Engineering Details": "9.3 Feature Engineering Details",
    "Extended Section 5: Machine Learning Models": "9.4 Machine Learning Models",
    "Extended Section 6: Training and Evaluation Flow": "9.5 Training and Evaluation Flow",
    "Extended Section 7: Prediction Logic": "9.6 Prediction Logic",
    "Extended Section 8: API Endpoint Details": "9.7 API Endpoint Details",
    "Extended Section 9: User Interface Explanation": "9.8 User Interface Explanation",
    "Extended Section 10: Excel Reports": "9.9 Excel Reports",
    "Extended Section 11: Visualization Outputs": "9.10 Visualization Outputs",
    "Extended Section 12: Security and Validation": "9.11 Security and Validation",
    "Extended Section 13: Detailed Test Case Matrix": "9.12 Detailed Test Case Matrix",
    "Extended Section 14: Algorithmic Representation": "9.13 Algorithmic Representation",
    "Extended Section 15: Deployment Procedure": "9.14 Deployment Procedure",
    "Extended Section 16: Maintenance Plan": "9.15 Maintenance Plan",
    "Extended Section 17: Advantages of the Project": "9.16 Advantages of the Project",
    "Extended Section 18: Limitations in Detail": "9.17 Detailed Limitations",
    "Extended Section 19: Future Enhancement Details": "9.18 Future Enhancement Details",
    "Extended Section 20: Final Project Review": "9.19 Final Project Review",
}


def paragraph_text(element):
    if not isinstance(element, CT_P):
        return ""
    return "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))


def main():
    doc = Document(TEMP)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in RENAME:
            paragraph.clear()
            run = paragraph.add_run(RENAME[text])
            run.bold = True
            run.font.name = "Times New Roman"
            if RENAME[text] == "Additional Project Documentation":
                paragraph.style = doc.styles["Heading 1"]
            else:
                paragraph.style = doc.styles["Heading 2"]

    body = doc._body._element
    children = list(body)

    references_idx = None
    added_start_idx = None
    for idx, child in enumerate(children):
        text = paragraph_text(child).strip()
        if text == "References / Bibliography":
            references_idx = idx
        if text == "Additional Project Documentation":
            added_start_idx = idx
            break

    if references_idx is None or added_start_idx is None:
        raise RuntimeError("Could not find references or additional documentation block.")

    # Include the page-break paragraph immediately before the added block when present.
    move_start = max(0, added_start_idx - 1)
    moving = children[move_start:]
    for element in moving:
        body.remove(element)

    children_after_remove = list(body)
    insert_before = None
    for child in children_after_remove:
        if paragraph_text(child).strip() == "References / Bibliography":
            insert_before = child
            break
    if insert_before is None:
        raise RuntimeError("Could not locate insertion point after removal.")

    for element in moving:
        body.insert(body.index(insert_before), element)

    doc.save(OUT)
    doc.save(COPY)
    print(OUT)
    print(COPY)


if __name__ == "__main__":
    main()
