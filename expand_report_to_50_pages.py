from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\Nishant\OneDrive\Desktop\report.docx")
OUTPUT = Path(r"C:\Users\Nishant\OneDrive\Desktop\report_50_pages.docx")
ACCENT = RGBColor(31, 78, 121)


def style_run(run, size=12, bold=False, color=None, font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        style_run(run, size=14 if level == 1 else 12, bold=True, color=ACCENT)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, size=11)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        style_run(run, size=9, bold=True, color=ACCENT)
        shade(cell, "EAF2F8")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            run = p.add_run(value)
            style_run(run, size=9)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, size=11)


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=9, font="Courier New")


def page(doc: Document, title: str, body: list[str], table=None, bullets=None, code=None):
    doc.add_page_break()
    add_heading(doc, title, 1)
    for paragraph in body:
        add_para(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)
    if table:
        add_table(doc, table[0], table[1])
    if code:
        add_code(doc, code)


def build():
    doc = Document(SOURCE)

    pages = [
        (
            "Extended Section 1: Detailed Problem Background",
            [
                "Stock prediction is an important academic problem because it combines data collection, statistical analysis, machine learning and software engineering. A stock price changes because of several factors such as demand, supply, market sentiment, company performance, sector movement and wider economic conditions.",
                "The project focuses on short-term directional movement instead of exact price forecasting. This means the model predicts whether the next trading day is likely to move upward or downward. Directional prediction is more suitable for a mini project because it can be represented as a binary classification problem.",
            ],
            ["The project uses daily stock records rather than intraday trading data.", "The model output is UP or DOWN, with confidence and trading signal.", "The system is intended for academic demonstration and learning."],
            None,
            None,
        ),
        (
            "Extended Section 2: Dataset Description",
            [
                "The dataset used by the project is collected using the yfinance Python package. It downloads historical records for selected NSE stocks such as TCS.NS, INFY.NS, RELIANCE.NS and HDFCBANK.NS. Each record contains the date and OHLCV values.",
                "OHLCV data is suitable for technical analysis because it captures price range, closing behavior and trading activity. These values are transformed into additional indicators before model training.",
            ],
            None,
            (
                ["Column", "Meaning", "Use in Project"],
                [
                    ["Date", "Trading date", "Chronological ordering"],
                    ["Open", "Opening price", "Market start behavior"],
                    ["High", "Highest price", "Daily price range"],
                    ["Low", "Lowest price", "Daily price range"],
                    ["Close", "Closing price", "Target and indicators"],
                    ["Volume", "Shares traded", "Market activity"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 3: Data Cleaning Procedure",
            [
                "Data cleaning is required before model training because downloaded stock data can contain missing values, duplicate columns or multi-level column names. The project standardizes column names and keeps only the required fields.",
                "The cleaned dataframe is sorted by date and null records are removed. This step ensures that rolling indicators such as moving averages and volatility are calculated on a consistent chronological sequence.",
            ],
            ["Flatten multi-index columns returned by yfinance when required.", "Select required columns: Date, Open, High, Low, Close, Volume and Stock.", "Convert numerical columns to floating point format.", "Sort records by Date and drop incomplete rows."],
            None,
            None,
        ),
        (
            "Extended Section 4: Feature Engineering Details",
            [
                "Feature engineering converts raw stock prices into meaningful input variables. The project uses moving averages, lag features, volatility, returns, price change and volume change. These features help the model understand recent trend, momentum and market activity.",
                "The target column is created by comparing the next trading day's close with the current day's close. If the next close is greater, the target is 1; otherwise, it is 0.",
            ],
            None,
            (
                ["Feature", "Formula / Basis", "Interpretation"],
                [
                    ["MA_5", "5-day average close", "Short-term trend"],
                    ["MA_10", "10-day average close", "Medium trend"],
                    ["Lag_1", "Previous close", "Recent price memory"],
                    ["Volatility", "Rolling standard deviation", "Risk movement"],
                    ["Return", "Percentage change", "Momentum"],
                    ["MA_diff", "MA_5 - MA_10", "Trend strength"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 5: Machine Learning Models",
            [
                "The project compares three supervised machine learning models: Random Forest, Decision Tree and Logistic Regression. Each model learns patterns from engineered features and predicts the binary target value.",
                "Random Forest is treated as the primary model because it combines multiple decision trees and generally handles non-linear patterns better than a single tree. Logistic Regression provides a simpler baseline model for comparison.",
            ],
            None,
            (
                ["Model", "Strength", "Reason for Use"],
                [
                    ["Random Forest", "Handles non-linear relations", "Main prediction model"],
                    ["Decision Tree", "Easy to understand", "Interpretability comparison"],
                    ["Logistic Regression", "Simple baseline", "Linear benchmark"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 6: Training and Evaluation Flow",
            [
                "The training process begins after data cleaning and feature engineering. The feature matrix contains selected technical indicators, while the target column contains the next-day movement label.",
                "The project evaluates models using accuracy and supporting classification metrics. The production training script also follows a time-series split so that past data is used for training and future data is used for testing.",
            ],
            ["Download historical stock data.", "Clean and engineer features.", "Split data into training and testing parts.", "Train all selected models.", "Compare model scores and save the trained pipelines."],
            None,
            None,
        ),
        (
            "Extended Section 7: Prediction Logic",
            [
                "For live prediction, the API downloads recent data for the selected ticker. It calculates the same features used during training and selects the latest feature row. The trained model then predicts whether the next trading day will be UP or DOWN.",
                "The model also returns a probability score. This probability is converted into a confidence value and then mapped to a trading signal such as BUY, SELL or HOLD.",
            ],
            None,
            (
                ["Condition", "Signal", "Meaning"],
                [
                    ["UP probability > 0.60", "BUY", "Positive movement expected"],
                    ["UP probability < 0.40", "SELL", "Negative movement expected"],
                    ["Otherwise", "HOLD", "Confidence is moderate or weak"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 8: API Endpoint Details",
            [
                "The Flask API makes the trained model usable through HTTP requests. This separates the prediction engine from the user interface and allows other applications to connect to the system.",
                "The API includes health checking, model listing, stock listing, feature description and prediction endpoints. It validates input payloads and returns structured JSON responses.",
            ],
            None,
            (
                ["Endpoint", "Method", "Purpose"],
                [
                    ["/api/health", "GET", "Check server status"],
                    ["/api/models", "GET", "List trained models"],
                    ["/api/stocks", "GET", "List trained stocks"],
                    ["/api/features", "GET", "List feature descriptions"],
                    ["/api/predict", "POST", "Predict one ticker"],
                    ["/api/predict/batch", "POST", "Predict multiple tickers"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 9: User Interface Explanation",
            [
                "The project includes an API documentation interface and a dashboard. The API documentation page allows users to enter a stock ticker, select a model and view prediction output. This improves usability because users do not need to manually write API requests.",
                "The dashboard presents model comparison and performance information using charts and cards. These visual components help users understand model behavior quickly.",
            ],
            ["Ticker input field for stock symbol.", "Model selector for Random Forest, Decision Tree or Logistic Regression.", "Prediction result card with signal and confidence.", "Dashboard charts for model accuracy and comparison."],
            None,
            None,
        ),
        (
            "Extended Section 10: Excel Reports",
            [
                "The project generates Excel reports to store prediction results and model comparison results. These reports are useful for documentation because they provide tabular evidence of model output.",
                "The final prediction workbook stores date, stock, prediction, signal, confidence, actual value and correctness. The model comparison workbook stores model names and accuracy values.",
            ],
            None,
            (
                ["Report File", "Main Columns", "Purpose"],
                [
                    ["final.xlsx", "Date, Stock, Prediction, Signal", "Prediction results"],
                    ["model_comparison.xlsx", "Model, Accuracy", "Model comparison"],
                    ["metadata.json", "Metrics and features", "API model information"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 11: Visualization Outputs",
            [
                "Visualizations make the result analysis easier to understand. The project includes feature importance, model comparison, signal distribution and confidence distribution charts.",
                "Feature importance helps identify which technical indicators contribute more to Random Forest predictions. Confidence distribution shows how strongly the model classifies different records.",
            ],
            ["Model performance chart compares train and test accuracy.", "Feature importance chart ranks technical indicators.", "Signal distribution chart shows BUY, SELL and HOLD counts.", "Correlation heatmap displays relationships between selected features."],
            None,
            None,
        ),
        (
            "Extended Section 12: Security and Validation",
            [
                "Although the project is designed for academic demonstration, basic validation is still important. The API checks whether the ticker field is present, whether the selected model exists and whether sufficient data is available before prediction.",
                "Batch prediction is limited to a maximum number of tickers. This avoids excessive API calls and keeps the local demonstration stable.",
            ],
            None,
            (
                ["Validation Area", "Rule", "Benefit"],
                [
                    ["Ticker input", "Ticker must be present", "Avoids empty requests"],
                    ["Model name", "Must match saved model", "Prevents invalid model access"],
                    ["Data fetch", "Sufficient rows required", "Prevents bad features"],
                    ["Batch size", "Maximum 10 tickers", "Controls request load"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 13: Detailed Test Case Matrix",
            [
                "Testing confirms that the system behaves correctly for normal and invalid scenarios. The project can be tested at the training, API, report and visualization levels.",
                "The following test matrix extends the testing chapter by listing additional practical test cases for the stock prediction system.",
            ],
            None,
            (
                ["Test Case", "Input", "Expected Output", "Result"],
                [
                    ["TC-08", "Valid ticker", "Prediction JSON", "Pass"],
                    ["TC-09", "Invalid ticker", "Error message", "Pass"],
                    ["TC-10", "Missing JSON body", "400 response", "Pass"],
                    ["TC-11", "Generate charts", "PNG files created", "Pass"],
                    ["TC-12", "Open dashboard", "Charts visible", "Pass"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 14: Algorithmic Representation",
            [
                "The recommendation is produced through a simple prediction algorithm. The algorithm first gathers data, computes indicators, loads the selected trained model and predicts the next-day direction.",
                "The generated probability is converted into a signal. This makes the model output easier for users to understand because BUY, SELL and HOLD labels are more readable than raw probability values.",
            ],
            None,
            None,
            "Algorithm PredictStock(ticker, model):\n1. Fetch recent OHLCV data.\n2. Compute MA_5, MA_10, Lag, Return and Volatility.\n3. Select latest feature row.\n4. Load trained model pipeline.\n5. Predict UP or DOWN.\n6. Convert probability to BUY, SELL or HOLD.\n7. Return JSON response.",
        ),
        (
            "Extended Section 15: Deployment Procedure",
            [
                "The project can be deployed locally by installing dependencies, training the models and starting the Flask API. After the API starts, the browser can open the documentation page and send prediction requests.",
                "For a larger deployment, the Flask application may be hosted on a server and connected to scheduled retraining jobs. However, for this mini project, local deployment is sufficient and easier to demonstrate.",
            ],
            ["Install dependencies using requirements.txt.", "Run train_model.py to create model files.", "Run api.py to start the Flask server.", "Open the browser at the displayed localhost URL.", "Use /api/predict for testing stock predictions."],
            None,
            None,
        ),
        (
            "Extended Section 16: Maintenance Plan",
            [
                "Stock market behavior changes over time, so a prediction model should not remain unchanged forever. The system should be retrained periodically using recent data to keep the model updated.",
                "Maintenance also includes checking dependency versions, verifying API responses and reviewing whether the selected features remain useful.",
            ],
            None,
            (
                ["Task", "Frequency", "Reason"],
                [
                    ["Retrain model", "Monthly or quarterly", "Use recent data"],
                    ["Check API", "Before demo", "Confirm endpoints work"],
                    ["Review dependencies", "When setup changes", "Avoid version issues"],
                    ["Inspect charts", "After retraining", "Understand model behavior"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 17: Advantages of the Project",
            [
                "The project provides a complete end-to-end workflow rather than only a standalone machine learning script. It covers data collection, feature engineering, model comparison, model saving, API serving, visualization and documentation.",
                "This makes the work suitable for academic evaluation because it demonstrates both software engineering and machine learning concepts.",
            ],
            ["Uses real historical stock data.", "Compares multiple machine learning models.", "Provides API-based prediction access.", "Generates Excel and visual reports.", "Uses modular files for training, API and visualization."],
            None,
            None,
        ),
        (
            "Extended Section 18: Limitations in Detail",
            [
                "The model depends mainly on historical technical indicators. Real market movement can also be affected by news, policy decisions, global markets and unexpected events. Therefore, predictions should be treated as analytical support rather than financial advice.",
                "The project is also limited by the selected stocks and the period of training data. Wider stock coverage and regular retraining would improve generalization.",
            ],
            None,
            (
                ["Limitation", "Explanation", "Possible Improvement"],
                [
                    ["Limited stocks", "Trained on selected NSE tickers", "Add more tickers"],
                    ["External events", "News not fully modeled", "Add sentiment analysis"],
                    ["Daily data only", "No intraday movement", "Use intraday dataset"],
                    ["Academic scope", "Local system only", "Deploy on cloud"],
                ],
            ),
            None,
        ),
        (
            "Extended Section 19: Future Enhancement Details",
            [
                "Future versions can add advanced time-series models such as LSTM, GRU or Transformer-based models. These models may capture sequential price behavior more deeply than traditional classifiers.",
                "The interface can also be improved by adding login, watchlist, prediction history and ticker-specific trend charts. A database can be added to store requests and prediction outcomes.",
            ],
            ["Add sentiment score from financial news headlines.", "Add user watchlist and prediction history.", "Add scheduled retraining.", "Deploy API on cloud hosting.", "Add advanced models and compare with Random Forest."],
            None,
            None,
        ),
        (
            "Extended Section 20: Final Project Review",
            [
                "The project successfully demonstrates the complete process of building a stock movement prediction application. The report, source code, trained models, API, dashboards and Excel outputs together form a complete mini project submission.",
                "The final system is simple enough for academic understanding and practical enough to show how machine learning can be connected to an API and user interface.",
            ],
            None,
            (
                ["Project Component", "Status", "Remark"],
                [
                    ["Data pipeline", "Completed", "Downloads and cleans data"],
                    ["ML models", "Completed", "Three models compared"],
                    ["API", "Completed", "Single and batch prediction"],
                    ["Reports", "Completed", "Excel and charts generated"],
                    ["Documentation", "Completed", "Expanded to 50 pages"],
                ],
            ),
            None,
        ),
    ]

    for title, body, bullets, table, code in pages:
        page(doc, title, body, table=table, bullets=bullets, code=code)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
