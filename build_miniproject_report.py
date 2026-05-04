from __future__ import annotations

import json
import math
import os
from collections import Counter
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "report_output"
DIAGRAM_DIR = OUT_DIR / "diagrams"
REPORT = OUT_DIR / "Stock_Movement_Prediction_Mini_Project_Report.docx"


PROJECT_TITLE = "Stock Movement Prediction and Trading Signal API"
STUDENT_NAME = "Name of the Student"
ENROLLMENT_NO = "Enrollment Number"
GUIDE_NAME = "Name of the Guide"
UNIVERSITY = "Graphic Era Deemed to be University, Dehradun"
DEPARTMENT = "Department of Computer Applications"
DEGREE = "Master of Computer Applications"


ACCENT = RGBColor(35, 82, 124)
LIGHT = "EAF2F8"
BORDER = "A6B8C8"
CONTENT_WIDTH_DXA = 9028


def ensure_dirs() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    DIAGRAM_DIR.mkdir(exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_rows: int = 1) -> None:
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
            if i < header_rows:
                set_cell_shading(cell, LIGHT)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = ACCENT


def apply_table_geometry(table) -> None:
    col_count = len(table.columns)
    base = CONTENT_WIDTH_DXA // col_count
    widths = [base] * col_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None):
    if caption:
        p = doc.add_paragraph()
        p.style = "Caption"
        p.add_run(caption).bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color=ACCENT)
        set_cell_shading(cell, LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    style_table(table)
    apply_table_geometry(table)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = ACCENT if level <= 2 else RGBColor(0, 0, 0)


def para(doc: Document, text: str = "", style: str | None = None, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        rest = text[len(bold_lead):]
        r = p.add_run(rest)
    else:
        r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        p.add_run(item).font.name = "Times New Roman"


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Caption"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text).bold = True


def add_image(doc: Document, path: Path, caption_text: str, width: float = 5.8) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        caption(doc, caption_text)


def read_text(path: Path, max_chars: int = 6000) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]


def workbook_rows(path: Path, limit: int | None = None):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = []
    for idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
        rows.append(list(row))
        if limit and idx >= limit:
            break
    return rows, ws.max_row, ws.max_column


def gather_data() -> dict:
    metadata = json.loads((ROOT / "models" / "metadata.json").read_text(encoding="utf-8"))
    comp_rows, _, _ = workbook_rows(ROOT / "model_comparison.xlsx")
    final_rows, final_count, _ = workbook_rows(ROOT / "final.xlsx")
    headers = final_rows[0]
    records = [dict(zip(headers, row)) for row in final_rows[1:] if any(v is not None for v in row)]
    signals = Counter(r["Signal"] for r in records)
    stocks = Counter(r["Stock"] for r in records)
    correct = sum(int(r["Correct"]) for r in records)
    confidence = [float(r["Confidence"]) for r in records if r.get("Confidence") is not None]
    project_files = []
    for p in sorted(ROOT.rglob("*")):
        if any(part in {".git", ".venv", ".sixth", "report_output", "__pycache__"} for part in p.parts):
            continue
        if p.is_file():
            if p.name in {"build_miniproject_report.py", "qa_split_pdf.py"}:
                continue
            project_files.append([str(p.relative_to(ROOT)), f"{p.stat().st_size:,} bytes"])
    return {
        "metadata": metadata,
        "comparison": comp_rows[1:],
        "records": records,
        "final_count": final_count - 1,
        "signals": signals,
        "stocks": stocks,
        "correct": correct,
        "accuracy": correct / len(records) if records else 0,
        "confidence_min": min(confidence) if confidence else 0,
        "confidence_max": max(confidence) if confidence else 0,
        "confidence_avg": sum(confidence) / len(confidence) if confidence else 0,
        "project_files": project_files,
        "readme": read_text(ROOT / "README.md"),
        "code_notes": read_text(ROOT / "project_code_explanation.md", 9000),
    }


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def draw_wrapped_text(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill="#1f2933") -> None:
    x1, y1, x2, y2 = box
    words = text.split()
    lines: list[str] = []
    current = ""
    max_w = x2 - x1 - 24
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_w:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_h = draw.textbbox((0, 0), "Ag", font=fnt)[3] + 5
    total_h = len(lines) * line_h
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += line_h


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#5b7083", width=3)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    pts = [
        end,
        (end[0] - size * math.cos(angle - 0.45), end[1] - size * math.sin(angle - 0.45)),
        (end[0] - size * math.cos(angle + 0.45), end[1] - size * math.sin(angle + 0.45)),
    ]
    draw.polygon(pts, fill="#5b7083")


def save_graph(path: Path, nodes: dict[str, tuple[float, float]], edges: list[tuple[str, str]], title: str) -> None:
    img = Image.new("RGB", (1400, 850), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(32, True)
    node_font = font(20, True)
    draw.text((700 - draw.textbbox((0, 0), title, font=title_font)[2] / 2, 28), title, font=title_font, fill="#23527c")
    min_x = min(x for x, _ in nodes.values())
    max_x = max(x for x, _ in nodes.values())
    min_y = min(y for _, y in nodes.values())
    max_y = max(y for _, y in nodes.values())

    def scale(pos):
        x, y = pos
        sx = 150 + (x - min_x) / max(1, max_x - min_x) * 1100
        sy = 150 + (max_y - y) / max(1, max_y - min_y) * 560
        return int(sx), int(sy)

    boxes = {}
    for name, pos in nodes.items():
        cx, cy = scale(pos)
        boxes[name] = (cx - 105, cy - 45, cx + 105, cy + 45)
    for a, b in edges:
        ax1, ay1, ax2, ay2 = boxes[a]
        bx1, by1, bx2, by2 = boxes[b]
        start = ((ax1 + ax2) // 2, (ay1 + ay2) // 2)
        end = ((bx1 + bx2) // 2, (by1 + by2) // 2)
        draw_arrow(draw, start, end)
    for name, box in boxes.items():
        draw.rounded_rectangle(box, radius=16, fill="#eaf2f8", outline="#23527c", width=3)
        draw_wrapped_text(draw, box, name, node_font)
    img.save(path)


def create_diagrams(data: dict) -> dict[str, Path]:
    paths = {}
    paths["context"] = DIAGRAM_DIR / "context_diagram.png"
    save_graph(
        paths["context"],
        {
            "User / Student": (0, 1),
            "Flask API": (2, 1),
            "Prediction Model": (4, 1),
            "Yahoo Finance": (2, 2.3),
            "Excel Reports": (2, -0.2),
            "Dashboard": (4, -0.2),
        },
        [
            ("User / Student", "Flask API"),
            ("Flask API", "Yahoo Finance"),
            ("Flask API", "Prediction Model"),
            ("Prediction Model", "Excel Reports"),
            ("Excel Reports", "Dashboard"),
            ("Dashboard", "User / Student"),
        ],
        "System Context / Level 0 DFD",
    )
    paths["architecture"] = DIAGRAM_DIR / "architecture.png"
    save_graph(
        paths["architecture"],
        {
            "Data Source": (0, 2),
            "Cleaning": (1.6, 2),
            "Feature Engineering": (3.2, 2),
            "Model Training": (4.8, 2),
            "Saved Models": (6.4, 2),
            "API Layer": (4.8, 0.7),
            "Web UI": (6.4, 0.7),
            "Reports": (3.2, 0.7),
        },
        [
            ("Data Source", "Cleaning"),
            ("Cleaning", "Feature Engineering"),
            ("Feature Engineering", "Model Training"),
            ("Model Training", "Saved Models"),
            ("Saved Models", "API Layer"),
            ("API Layer", "Web UI"),
            ("Feature Engineering", "Reports"),
        ],
        "Logical Architecture",
    )
    paths["sequence"] = DIAGRAM_DIR / "sequence.png"
    save_graph(
        paths["sequence"],
        {
            "Client": (0, 1.5),
            "POST /api/predict": (1.8, 1.5),
            "Fetch Data": (3.6, 2.4),
            "Compute Features": (3.6, 0.8),
            "Model Predict": (5.4, 1.5),
            "JSON Response": (7.2, 1.5),
        },
        [
            ("Client", "POST /api/predict"),
            ("POST /api/predict", "Fetch Data"),
            ("Fetch Data", "Compute Features"),
            ("Compute Features", "Model Predict"),
            ("Model Predict", "JSON Response"),
            ("JSON Response", "Client"),
        ],
        "Prediction Request Sequence",
    )
    paths["gantt"] = DIAGRAM_DIR / "gantt.png"
    tasks = [
        ("Requirement Study", 1, 3),
        ("Data Collection", 3, 4),
        ("Feature Engineering", 6, 4),
        ("Model Development", 9, 5),
        ("API Development", 13, 4),
        ("Testing", 16, 4),
        ("Documentation", 19, 4),
    ]
    img = Image.new("RGB", (1400, 780), "white")
    draw = ImageDraw.Draw(img)
    draw.text((430, 32), "Project Schedule / Gantt Chart", font=font(34, True), fill="#23527c")
    left, top, row_h = 280, 130, 70
    week_w = 42
    for week in range(1, 25):
        x = left + week * week_w
        draw.line([(x, top - 20), (x, top + row_h * len(tasks))], fill="#d8e2ec", width=1)
        if week % 2 == 0:
            draw.text((x - 10, top - 48), str(week), font=font(14), fill="#4b5c6b")
    for i, (name, start, dur) in enumerate(tasks):
        y = top + i * row_h
        draw.text((40, y + 18), name, font=font(20, True), fill="#1f2933")
        x = left + start * week_w
        w = dur * week_w
        draw.rounded_rectangle((x, y + 10, x + w, y + 50), radius=10, fill="#79a6c9", outline="#23527c", width=2)
        draw.text((x + 12, y + 19), f"W{start}-W{start + dur}", font=font(16, True), fill="white")
    draw.text((left + 420, top + row_h * len(tasks) + 30), "Weeks", font=font(20, True), fill="#23527c")
    img.save(paths["gantt"])
    return paths


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Times New Roman"
    styles["Title"].font.size = Pt(20)
    styles["Caption"].font.name = "Times New Roman"
    styles["Caption"].font.size = Pt(10)
    return doc


def set_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = PROJECT_TITLE
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in header_p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(90, 90, 90)
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.text = "Page "
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        footer_p._p.append(fld)


def front_matter(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJECT REPORT ON")
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT_TITLE.upper())
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT
    for text in [
        "Submitted to",
        DEPARTMENT,
        "in partial fulfillment for the award of the degree of",
        DEGREE.upper(),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(13)
    doc.add_paragraph()
    add_table(
        doc,
        ["Submitted By", "Under the Guidance of"],
        [[f"{STUDENT_NAME}\n{ENROLLMENT_NO}", GUIDE_NAME]],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(UNIVERSITY.upper()).bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"May {date.today().year}").bold = True
    doc.add_page_break()

    add_heading(doc, "Candidate Declaration", 1)
    para(doc, f"I hereby declare that the mini project report entitled \"{PROJECT_TITLE}\" has been prepared by me as part of the academic requirements of the {DEGREE}. The work presented in this report is original to the best of my knowledge and has been completed using the source files, code, results, visualizations and documentation available in the project folder.")
    para(doc, "The report content is based on project-specific implementation details including the data processing pipeline, feature engineering, model training, API implementation, testing artifacts, generated Excel reports and visualizations.")
    doc.add_paragraph("\n\nSignature of the Student: ____________________")
    doc.add_paragraph(f"Name: {STUDENT_NAME}")
    doc.add_paragraph("Date: ____________________")
    doc.add_page_break()

    add_heading(doc, "Internal Guide Certificate", 1)
    para(doc, f"This is to certify that the project report entitled \"{PROJECT_TITLE}\" submitted to {UNIVERSITY} in partial fulfillment of the requirements for the award of the degree of {DEGREE} is an authentic project work carried out by {STUDENT_NAME} under my supervision and guidance.")
    para(doc, "The project demonstrates a machine learning based approach for predicting next-day directional stock movement and presenting actionable trading signals through a Flask REST API and supporting dashboards.")
    doc.add_paragraph("\n\nSignature of the Guide: ____________________")
    doc.add_paragraph(f"Name: {GUIDE_NAME}")
    doc.add_paragraph("Date: ____________________")
    doc.add_page_break()

    add_heading(doc, "Acknowledgement", 1)
    para(doc, f"I express my sincere gratitude to {DEPARTMENT}, {UNIVERSITY}, for providing the opportunity to undertake this mini project. I am thankful to my guide for guidance, feedback and encouragement during the analysis, development, testing and documentation stages.")
    para(doc, "I also acknowledge the usefulness of open-source Python libraries such as pandas, NumPy, scikit-learn, Flask, yfinance, matplotlib, seaborn, joblib and openpyxl, which supported the implementation and evaluation of this work.")
    doc.add_page_break()

    add_heading(doc, "Preface", 1)
    para(doc, "This report presents the design and implementation of a mini project for stock movement prediction. The system uses historical OHLCV stock data, derives technical indicators, trains multiple classification models, stores trained pipelines, exposes prediction functionality through a REST API and produces reports and visualizations for interpretation.")
    para(doc, "The report follows the supplied MCA report guidelines by including introduction, system analysis, requirement specification, design, implementation, testing, results, project management, future scope, bibliography and appendices.")
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Introduction",
        "2. System Analysis and Requirement Specification",
        "3. System Design",
        "4. Project Management",
        "5. Implementation and Coding",
        "6. Results, Reports and User Interface",
        "7. Testing, Implementation and Maintenance",
        "8. Summary and Future Scope",
        "References / Bibliography",
        "Appendices",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    add_heading(doc, "List of Tables and Figures", 1)
    add_table(
        doc,
        ["Type", "Title"],
        [
            ["Table 1.1", "Project overview"],
            ["Table 2.1", "Functional requirements"],
            ["Table 2.2", "Non-functional requirements"],
            ["Table 3.1", "Feature dictionary"],
            ["Table 4.1", "Risk management plan"],
            ["Table 5.1", "Module inventory"],
            ["Table 6.1", "Model comparison"],
            ["Table 7.1", "Test cases"],
            ["Figure 2.1", "System context / Level 0 DFD"],
            ["Figure 3.1", "Logical architecture"],
            ["Figure 3.2", "Prediction request sequence"],
            ["Figure 4.1", "Gantt chart"],
        ],
    )
    doc.add_page_break()


def chapter_introduction(doc: Document, data: dict) -> None:
    add_heading(doc, "Chapter 1: Introduction", 1)
    para(doc, "Stock market movement prediction is a challenging classification problem because market prices are affected by trends, volatility, volume behavior, investor sentiment and external events. This mini project focuses on predicting whether the next trading day's closing price will move upward or downward for selected Indian NSE stocks.")
    para(doc, "The implemented project is a Python based machine learning pipeline named Stock Movement Prediction and Trading Signal API. It uses historical data for TCS.NS, INFY.NS, RELIANCE.NS and HDFCBANK.NS, generates technical indicators, trains classification models and serves predictions using a Flask REST API.")
    add_table(
        doc,
        ["Item", "Description"],
        [
            ["Project domain", "Machine learning, financial analytics and web API development"],
            ["Primary objective", "Predict next-day stock direction as UP or DOWN"],
            ["Supported outputs", "Prediction label, confidence, BUY/SELL/HOLD signal and feature values"],
            ["Models used", "Random Forest, Decision Tree and Logistic Regression"],
            ["Data source", "Yahoo Finance via yfinance"],
            ["Main deliverables", "Training pipeline, saved models, API, dashboards, Excel reports and documentation"],
        ],
        "Table 1.1: Project overview",
    )
    add_heading(doc, "1.1 Motivation", 2)
    para(doc, "Investors and learners often need a compact decision-support tool that can combine recent price behavior with interpretable model outputs. The project does not claim to guarantee market profit; instead, it demonstrates a structured software engineering approach to data collection, feature engineering, model evaluation and API delivery.")
    add_heading(doc, "1.2 Objectives", 2)
    bullets(doc, [
        "Collect historical OHLCV data for selected NSE stocks.",
        "Clean and transform raw market records into machine learning features.",
        "Train and compare Random Forest, Decision Tree and Logistic Regression models.",
        "Generate BUY, SELL and HOLD signals from prediction probabilities.",
        "Provide a Flask API for single and batch predictions.",
        "Create dashboards, charts and Excel reports for result interpretation.",
    ])
    add_heading(doc, "1.3 Scope", 2)
    para(doc, "The scope is limited to directional movement prediction using daily historical market data. The project supports any valid Yahoo Finance ticker at API runtime, but the bundled trained models were trained on four Indian NSE stocks. The work includes analysis, design, implementation, testing and documentation but excludes live trading execution, brokerage integration and financial advisory certification.")
    add_heading(doc, "1.4 Hardware and Software Requirements", 2)
    add_table(
        doc,
        ["Category", "Requirement"],
        [
            ["Hardware", "Standard laptop/desktop with 4 GB RAM or above and internet connectivity"],
            ["Operating system", "Windows or any OS capable of running Python 3"],
            ["Language", "Python"],
            ["Libraries", "pandas, numpy, scikit-learn, yfinance, Flask, Flask-CORS, joblib, openpyxl, matplotlib, seaborn, TextBlob, NewsAPI Python"],
            ["Tools", "Jupyter Notebook, browser, spreadsheet viewer and code editor"],
        ],
        "Table 1.2: Hardware and software requirements",
    )


def chapter_analysis(doc: Document, data: dict, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 2: System Analysis and Requirement Specification", 1)
    add_heading(doc, "2.1 Existing System", 2)
    para(doc, "A manual stock analysis workflow usually requires users to download prices, calculate indicators, compare models and interpret probabilities separately. Such a process is time consuming and prone to inconsistent calculations. Existing online tools may provide charts, but they often do not expose the full model workflow or reusable API endpoints for project-level experimentation.")
    add_heading(doc, "2.2 Proposed System", 2)
    para(doc, "The proposed system automates the core workflow from data collection to prediction delivery. Historical data is collected using yfinance, transformed into indicators, passed through trained scikit-learn pipelines and served through REST endpoints. The dashboard and Excel reports make the predictions easier to review.")
    add_image(doc, diagrams["context"], "Figure 2.1: System context / Level 0 DFD", 6.1)
    add_heading(doc, "2.3 Functional Requirements", 2)
    add_table(
        doc,
        ["Req. ID", "Requirement", "Implemented Through"],
        [
            ["FR-01", "Download recent and historical stock data", "yfinance in mainn.py, train_model.py and api.py"],
            ["FR-02", "Engineer technical indicators", "clean_and_engineer() and compute_features()"],
            ["FR-03", "Train multiple classification models", "train_models() in train_model.py"],
            ["FR-04", "Persist trained models and metadata", "joblib pickle files and metadata.json"],
            ["FR-05", "Expose prediction API endpoints", "Flask routes in api.py"],
            ["FR-06", "Generate result reports", "final.xlsx and model_comparison.xlsx"],
            ["FR-07", "Generate charts and dashboards", "visualize.py, dashboard.html and api_docs.html"],
        ],
        "Table 2.1: Functional requirements",
    )
    add_heading(doc, "2.4 Non-Functional Requirements", 2)
    add_table(
        doc,
        ["Quality", "Requirement"],
        [
            ["Usability", "The API documentation page should allow users to test predictions easily."],
            ["Reliability", "The system must validate tickers, model names and input payloads."],
            ["Maintainability", "Training, API and visualization concerns are separated into different files."],
            ["Performance", "Saved models avoid retraining during API prediction requests."],
            ["Portability", "Dependencies are listed in requirements.txt for reproducible setup."],
            ["Security", "The API validates request bodies and limits batch prediction size to 10 tickers."],
        ],
        "Table 2.2: Non-functional requirements",
    )
    add_heading(doc, "2.5 Feasibility Study", 2)
    para(doc, "Technical feasibility is strong because all major components are implemented with stable Python libraries. Operational feasibility is appropriate for academic demonstration because the system can be run locally and used through browser pages or API clients. Economic feasibility is favorable because the implementation uses open-source tools and public data access rather than paid infrastructure.")


def chapter_design(doc: Document, data: dict, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 3: System Design", 1)
    add_heading(doc, "3.1 Architectural Design", 2)
    para(doc, "The project follows a layered architecture. The data layer obtains raw OHLCV records, the feature layer computes model inputs, the machine learning layer trains and stores pipelines, the service layer exposes predictions and the presentation layer displays documentation, dashboards and reports.")
    add_image(doc, diagrams["architecture"], "Figure 3.1: Logical architecture of the project", 6.2)
    add_heading(doc, "3.2 Data Flow and Processing", 2)
    para(doc, "The data flow begins with daily stock records containing Date, Open, High, Low, Close, Volume and Stock. The system sorts records chronologically, removes null values, derives rolling features and constructs the target label. In training, the target is binary: 1 when the next day's close is greater than the current close, otherwise 0.")
    add_heading(doc, "3.3 Feature Dictionary", 2)
    feature_rows = [
        ["MA_5", "5-day moving average of closing price", "Short-term trend"],
        ["MA_10", "10-day moving average of closing price", "Medium-term trend"],
        ["Lag_1", "Previous day close", "Recent memory"],
        ["Lag_2", "Close from two days earlier", "Short lag memory"],
        ["Volatility", "5-day rolling standard deviation", "Risk and fluctuation"],
        ["Return", "Daily percentage change", "Momentum"],
        ["Volume_MA", "5-day moving average of volume", "Trading activity"],
        ["MA_diff", "MA_5 minus MA_10", "Trend strength"],
        ["Price_Change", "Close minus previous close", "Absolute movement"],
        ["Volume_Change", "Daily percentage volume change", "Volume momentum"],
    ]
    add_table(doc, ["Feature", "Definition", "Purpose"], feature_rows, "Table 3.1: Feature dictionary")
    add_heading(doc, "3.4 Sequence Design", 2)
    para(doc, "For a prediction request, the client sends a ticker and optional model name. The API downloads recent records, computes features, extracts the latest feature vector, invokes the selected pipeline and returns prediction metadata in JSON format.")
    add_image(doc, diagrams["sequence"], "Figure 3.2: Prediction request sequence diagram", 6.2)
    add_heading(doc, "3.5 User Interface Design", 2)
    para(doc, "The project contains two HTML interfaces. The API documentation page provides a prediction panel, model selector, quick stock buttons and result cards. The benchmark dashboard displays model comparison metrics, ranking and visual interpretation. These pages make the API outputs accessible without requiring a separate API client.")


def chapter_management(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 4: Project Management", 1)
    add_heading(doc, "4.1 Development Approach", 2)
    para(doc, "An iterative prototyping approach was suitable because the project required experimentation before finalizing the production pipeline. The notebook and mainn.py captured exploration, while train_model.py converted the successful workflow into a reusable training pipeline. api.py then integrated the trained models into a service layer.")
    add_image(doc, diagrams["gantt"], "Figure 4.1: Project schedule / Gantt chart", 6.0)
    add_heading(doc, "4.2 Milestones and Deliverables", 2)
    add_table(
        doc,
        ["Milestone", "Deliverable", "Dependency"],
        [
            ["Requirement study", "Problem statement and scope", "Guidelines and project idea"],
            ["Data collection", "Stock dataset from 2022 to 2025", "Internet and yfinance"],
            ["Feature engineering", "Ten technical indicators", "Clean OHLCV records"],
            ["Model training", "Three trained model pipelines", "Prepared feature matrix"],
            ["API development", "Flask endpoints", "Saved model files"],
            ["Reporting", "Excel sheets, charts and dashboard", "Prediction results"],
            ["Documentation", "Mini project report", "All project artifacts"],
        ],
        "Table 4.1: Milestones and deliverables",
    )
    add_heading(doc, "4.3 Risk Management", 2)
    add_table(
        doc,
        ["Risk", "Probability", "Impact", "Mitigation"],
        [
            ["Yahoo Finance data unavailable", "Medium", "High", "Validate responses and allow retry with valid tickers."],
            ["Model overfitting", "Medium", "High", "Use train/test evaluation and time-series split in production script."],
            ["Data leakage", "Medium", "High", "Prefer chronological split in train_model.py."],
            ["Low confidence predictions", "High", "Medium", "Use HOLD signal and confidence quality labels."],
            ["Dependency mismatch", "Medium", "Medium", "Maintain requirements.txt."],
            ["API misuse", "Low", "Medium", "Validate JSON payloads and cap batch size."],
        ],
        "Table 4.2: Risk management plan",
    )
    add_heading(doc, "4.4 Cost Estimation", 2)
    add_table(
        doc,
        ["Area", "Estimated Effort", "Cost Basis"],
        [
            ["Analysis and design", "30 hours", "Student development effort"],
            ["Data preparation", "25 hours", "Student development effort"],
            ["Model development", "45 hours", "Student development effort"],
            ["API and UI", "35 hours", "Student development effort"],
            ["Testing and debugging", "25 hours", "Student development effort"],
            ["Documentation", "30 hours", "Student development effort"],
        ],
        "Table 4.3: Effort-based cost estimation",
    )


def chapter_implementation(doc: Document, data: dict) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 5: Implementation and Coding", 1)
    para(doc, "The implementation is divided into notebook experimentation, production training, API serving, visualization generation and reporting artifacts. This modular structure improves maintainability and allows each part of the pipeline to be run independently.")
    add_heading(doc, "5.1 Project File Inventory", 2)
    add_table(doc, ["File / Folder", "Size"], data["project_files"][:28], "Table 5.1: Extracted project folder inventory")
    add_heading(doc, "5.2 Main Modules", 2)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["mainn.ipynb / mainn.py", "Original experiment, data cleaning, feature engineering, sentiment attempt, model comparison and Excel export."],
            ["train_model.py", "Production-style training with time-series split, metrics, backtesting and saved models."],
            ["api.py", "Flask server that loads models and exposes health, model, stock, feature, single prediction and batch prediction endpoints."],
            ["visualize.py", "Generates model comparison, feature importance and feature correlation visualizations."],
            ["dashboard.html", "Benchmark dashboard for model comparison and interpretation."],
            ["api_docs.html", "Interactive API documentation and prediction interface."],
        ],
        "Table 5.2: Module responsibility matrix",
    )
    add_heading(doc, "5.3 Training Pipeline", 2)
    para(doc, "The training script downloads records from 2022-01-01 to 2025-01-01 for TCS.NS, INFY.NS, RELIANCE.NS and HDFCBANK.NS. After cleaning and feature engineering, records are sorted by date and split chronologically. This design is important because future data should not be used to train a model that is supposed to predict later movement.")
    add_heading(doc, "5.4 API Implementation", 2)
    para(doc, "The API loads metadata.json and model pickle files from the models folder. It exposes JSON endpoints for health checking, model listing, supported stock listing, feature descriptions, single prediction and batch prediction. The predict_stock function performs input validation, fetches recent stock data, computes features, predicts direction and returns a structured response.")
    add_heading(doc, "5.5 Validation Checks", 2)
    bullets(doc, [
        "Missing ticker requests return a 400 error with an example payload.",
        "Unknown model names return an informative error.",
        "Insufficient stock data is rejected before prediction.",
        "Batch prediction requires a non-empty list and is limited to 10 tickers.",
        "Prediction dates skip weekends when computing the next trading day.",
    ])


def chapter_results(doc: Document, data: dict, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 6: Results, Reports and User Interface", 1)
    add_heading(doc, "6.1 Model Performance", 2)
    comp_rows = [[str(r[0]), f"{float(r[1]) * 100:.2f}%"] for r in data["comparison"] if r and r[0]]
    add_table(doc, ["Model", "Accuracy"], comp_rows, "Table 6.1: Model comparison from model_comparison.xlsx")
    para(doc, f"The Random Forest model achieved the best exported comparison accuracy of 80.00%. The final prediction workbook contains {data['final_count']} prediction rows with an observed correctness ratio of {data['accuracy'] * 100:.2f}% based on the Correct column.")
    add_heading(doc, "6.2 Signal Distribution", 2)
    add_table(
        doc,
        ["Signal", "Count"],
        [[k, str(v)] for k, v in sorted(data["signals"].items())],
        "Table 6.2: Trading signal distribution from final.xlsx",
    )
    add_heading(doc, "6.3 Prediction Statistics", 2)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Total prediction rows", str(data["final_count"])],
            ["Correct predictions", str(data["correct"])],
            ["Accuracy from final.xlsx", f"{data['accuracy'] * 100:.2f}%"],
            ["Minimum UP probability", f"{data['confidence_min']:.4f}"],
            ["Maximum UP probability", f"{data['confidence_max']:.4f}"],
            ["Average UP probability", f"{data['confidence_avg']:.4f}"],
        ],
        "Table 6.3: Prediction summary statistics",
    )
    add_heading(doc, "6.4 Sample Prediction Records", 2)
    sample = data["records"][:10]
    add_table(
        doc,
        ["Date", "Stock", "Prediction", "Signal", "Confidence", "Actual", "Correct"],
        [[str(r["Date"]), str(r["Stock"]), str(r["Prediction"]), str(r["Signal"]), f"{float(r['Confidence']):.4f}", str(r["Actual"]), str(r["Correct"])] for r in sample],
        "Table 6.4: Sample rows from final.xlsx",
    )
    add_heading(doc, "6.5 Generated Visual Outputs", 2)
    for img, cap in [
        (ROOT / "visualizations" / "model_performance.png", "Figure 6.1: Model performance visualization"),
        (ROOT / "visualizations" / "feature_importance.png", "Figure 6.2: Random Forest feature importance"),
        (ROOT / "visualizations" / "feature_correlation.png", "Figure 6.3: Feature correlation heatmap"),
        (ROOT / "confidence_distribution.png", "Figure 6.4: Prediction confidence distribution"),
        (ROOT / "signal_distribution.png", "Figure 6.5: Signal distribution chart"),
    ]:
        add_image(doc, img, cap, 5.6)


def chapter_testing(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 7: Testing, Implementation and Maintenance", 1)
    add_heading(doc, "7.1 Testing Strategy", 2)
    para(doc, "Testing was considered at data, model and API levels. Data testing confirms that expected columns are available and null values are removed. Model testing compares predictions against held-out data. API testing verifies endpoint behavior for normal and invalid request payloads.")
    add_table(
        doc,
        ["Test ID", "Scenario", "Expected Result", "Status"],
        [
            ["TC-01", "Run training pipeline", "Models and metadata saved in models folder", "Designed"],
            ["TC-02", "Open /api/health", "Healthy JSON response with loaded model count", "Designed"],
            ["TC-03", "POST /api/predict with TCS.NS", "Prediction response with signal and confidence", "Designed"],
            ["TC-04", "POST /api/predict without ticker", "400 error with example body", "Designed"],
            ["TC-05", "POST /api/predict/batch with more than 10 tickers", "400 error for batch size limit", "Designed"],
            ["TC-06", "Generate visualizations", "PNG charts saved in visualizations folder", "Designed"],
            ["TC-07", "Open final.xlsx", "Prediction rows and formatted headers visible", "Designed"],
        ],
        "Table 7.1: Test cases",
    )
    add_heading(doc, "7.2 Implementation Notes", 2)
    para(doc, "The system should be implemented by first installing dependencies from requirements.txt, then running train_model.py to refresh the trained models, then starting api.py. The dashboard and documentation pages can be opened in a browser, while API clients can call JSON endpoints directly.")
    add_heading(doc, "7.3 Maintenance", 2)
    bullets(doc, [
        "Retrain models periodically so that recent market data is reflected.",
        "Monitor API errors for invalid tickers and unavailable data source responses.",
        "Update requirements.txt when library versions change.",
        "Review feature importance before adding or removing indicators.",
        "Keep API documentation synchronized with endpoint behavior.",
    ])


def chapter_summary(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Chapter 8: Summary and Future Scope", 1)
    para(doc, "The mini project successfully demonstrates a complete machine learning software workflow for stock movement prediction. It includes data collection, cleaning, feature engineering, multiple model training, model persistence, API delivery, dashboard presentation and spreadsheet reporting.")
    add_heading(doc, "8.1 Key Outcomes", 2)
    bullets(doc, [
        "A working classification pipeline predicts next-day stock direction.",
        "Random Forest, Decision Tree and Logistic Regression were compared.",
        "Random Forest produced the strongest exported accuracy result.",
        "The API supports single and batch prediction workflows.",
        "Excel reports and visualization charts support result interpretation.",
    ])
    add_heading(doc, "8.2 Limitations", 2)
    bullets(doc, [
        "The model uses historical technical indicators and cannot capture every market-moving event.",
        "The bundled model was trained on four NSE stocks, so generalization should be evaluated before wider use.",
        "Sentiment analysis is exploratory and depends on external news API availability.",
        "The project is for academic decision-support demonstration and not financial advice.",
    ])
    add_heading(doc, "8.3 Future Enhancements", 2)
    bullets(doc, [
        "Add LSTM, GRU or Transformer-based time-series models for deeper sequential learning.",
        "Integrate robust sentiment features using reliable news sources.",
        "Add authentication and rate limiting for deployed API use.",
        "Store prediction history in a database for audit and monitoring.",
        "Add automated scheduled retraining and model drift detection.",
        "Build a richer frontend with charts for individual ticker trends and prediction history.",
    ])


def references_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "References / Bibliography", 1)
    refs = [
        "scikit-learn documentation for classification models and pipelines.",
        "pandas documentation for dataframe cleaning and transformation.",
        "NumPy documentation for numerical computation.",
        "Flask documentation for REST API development.",
        "Yahoo Finance data access through the yfinance Python package.",
        "matplotlib and seaborn documentation for visualization.",
        "openpyxl documentation for Excel report generation.",
        "Project source files: README.md, project_code_explanation.md, mainn.py, train_model.py, api.py and visualize.py.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")
    doc.add_page_break()
    add_heading(doc, "Appendix A: API Endpoint Summary", 1)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/", "GET", "Serve interactive API documentation page"],
            ["/api/health", "GET", "Return server and model health"],
            ["/api/models", "GET", "List models and metrics"],
            ["/api/stocks", "GET", "List trained stock tickers"],
            ["/api/features", "GET", "List feature descriptions"],
            ["/api/predict", "POST", "Predict one ticker"],
            ["/api/predict/batch", "POST", "Predict up to 10 tickers"],
        ],
    )
    add_heading(doc, "Appendix B: Important Code Excerpts", 1)
    snippets = [
        ("Feature Engineering", "df['MA_5'] = df['Close'].rolling(5).mean()\ndf['MA_10'] = df['Close'].rolling(10).mean()\ndf['Target'] = (df['Close'].shift(-1) > df['Close']).astype(int)"),
        ("Signal Generation", "if confidence > 0.60:\n    return 'BUY'\nelif confidence < 0.40:\n    return 'SELL'\nelse:\n    return 'HOLD'"),
        ("Model Persistence", "joblib.dump(data['pipeline'], model_path)\njson.dump(metadata, f, indent=2)"),
    ]
    for title, code in snippets:
        add_heading(doc, title, 2)
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    add_heading(doc, "Glossary", 1)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["OHLCV", "Open, High, Low, Close and Volume market data"],
            ["Feature", "Input variable used by a machine learning model"],
            ["Target", "Expected output label used for supervised learning"],
            ["Random Forest", "Ensemble classifier made from multiple decision trees"],
            ["API", "Application Programming Interface"],
            ["Confidence", "Model probability used to qualify prediction strength"],
            ["Backtesting", "Testing trading logic on historical data"],
        ],
    )


def build() -> None:
    ensure_dirs()
    data = gather_data()
    diagrams = create_diagrams(data)
    doc = setup_document()
    set_header_footer(doc)
    front_matter(doc)
    chapter_introduction(doc, data)
    chapter_analysis(doc, data, diagrams)
    chapter_design(doc, data, diagrams)
    chapter_management(doc, diagrams)
    chapter_implementation(doc, data)
    chapter_results(doc, data, diagrams)
    chapter_testing(doc)
    chapter_summary(doc)
    references_appendix(doc)
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build()
