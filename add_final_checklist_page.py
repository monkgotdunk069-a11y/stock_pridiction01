from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PATH = Path(r"C:\Users\Nishant\OneDrive\Desktop\report.docx")
COPY = Path(r"C:\Users\Nishant\OneDrive\Desktop\report_arranged_50_pages.docx")
ACCENT = RGBColor(31, 78, 121)


def style_run(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_block() -> Document:
    block = Document()
    block.add_page_break()
    h = block.add_heading("9.20 Documentation Checklist", level=2)
    for run in h.runs:
        style_run(run, 12, True, ACCENT)
    p = block.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        "This checklist summarizes how the submitted report satisfies the expected project documentation areas. "
        "It is included to make the report easier to verify against the project guidelines and to connect the "
        "implementation work with the required academic sections."
    )
    style_run(r)

    table = block.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Guideline Area", "Included Content", "Status"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(header)
        style_run(rr, 9, True, ACCENT)
        shade(cell, "EAF2F8")
    rows = [
        ["Introduction and Objectives", "Problem background, motivation, objectives and scope", "Completed"],
        ["System Analysis", "Existing system, proposed system, requirements and feasibility", "Completed"],
        ["System Design", "Architecture, data flow, sequence design and feature dictionary", "Completed"],
        ["Implementation", "Data collection, preprocessing, model training, API and dashboard", "Completed"],
        ["Testing", "Test strategy, test cases, validation checks and expected results", "Completed"],
        ["Reports and Results", "Model comparison, Excel outputs, charts and prediction samples", "Completed"],
        ["Future Scope", "Limitations, enhancements and maintenance plan", "Completed"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            rr = cells[idx].paragraphs[0].add_run(value)
            style_run(rr, 9)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return block


def main():
    doc = Document(PATH)
    if any("Extended Section" in p.text for p in doc.paragraphs):
        raise RuntimeError("Extended Section text still exists; clean headings first.")

    ref_para = next((p for p in doc.paragraphs if p.text.strip() == "References / Bibliography"), None)
    if ref_para is None:
        raise RuntimeError("References / Bibliography not found.")

    block = build_block()
    block_elements = list(block._body._element)[:-1]  # skip sectPr from temp doc
    body = doc._body._element
    ref_el = ref_para._p
    for element in block_elements:
        body.insert(body.index(ref_el), deepcopy(element))

    doc.save(PATH)
    doc.save(COPY)
    print(PATH)


if __name__ == "__main__":
    main()
